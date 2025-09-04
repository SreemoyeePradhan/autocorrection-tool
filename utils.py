# utils.py
import re
from typing import List, Dict, Tuple, Any
from difflib import SequenceMatcher
from io import BytesIO

# Optional dependencies
try:
    from docx import Document
except Exception:
    Document = None

try:
    from gtts import gTTS
except Exception:
    gTTS = None

# ---------- Preprocessing & Code Block Handling ----------

CODE_BLOCK_PATTERN = re.compile(r"```.*?```", re.DOTALL)

def preprocess_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())

def mask_code_blocks(text: str) -> Tuple[str, Dict[str, str]]:
    """
    Replace ```code``` blocks with placeholders to avoid editing them.
    Returns masked_text and a dict placeholder->original_block.
    """
    blocks = {}
    def repl(match):
        key = f"[[[CODE_BLOCK_{len(blocks)}]]]"
        blocks[key] = match.group(0)
        return key
    masked = CODE_BLOCK_PATTERN.sub(repl, text)
    return masked, blocks

def unmask_code_blocks(text: str, blocks: Dict[str, str]) -> str:
    for k, v in blocks.items():
        text = text.replace(k, v)
    return text

# ---------- Diff & Inline Edit Helpers ----------

def highlight_changes(original: str, corrected: str) -> str:
    """
    Token-level diff with color:
      - equal: normal
      - replace: red new token
      - insert: blue new token
      - delete: gray strikethrough old token
    """
    orig_tokens = original.split()
    corr_tokens = corrected.split()
    matcher = SequenceMatcher(None, orig_tokens, corr_tokens)
    out = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            out.extend(orig_tokens[i1:i2])
        elif tag == "replace":
            out.extend(f"<span style='color:red'>{t}</span>" for t in corr_tokens[j1:j2])
        elif tag == "insert":
            out.extend(f"<span style='color:blue'>{t}</span>" for t in corr_tokens[j1:j2])
        elif tag == "delete":
            out.extend(f"<span style='color:gray;text-decoration:line-through'>{t}</span>" for t in orig_tokens[i1:i2])
    return " ".join(out)

def diff_pairs(original: str, corrected: str) -> List[Dict[str, Any]]:
    """
    Produce a structured list of edits for inline accept/reject UI.
    Each item: {type, before_tokens, after_tokens, i1, i2, j1, j2}
    """
    o = original.split()
    c = corrected.split()
    ops = SequenceMatcher(None, o, c).get_opcodes()
    edits = []
    for tag, i1, i2, j1, j2 in ops:
        if tag == "equal":
            continue
        edits.append({
            "type": tag, "before_tokens": o[i1:i2], "after_tokens": c[j1:j2],
            "i1": i1, "i2": i2, "j1": j1, "j2": j2
        })
    return edits

def apply_selected_edits(original: str, corrected: str, accept_mask: List[bool]) -> str:
    """
    Apply only selected replacements/inserts; reject others by keeping original.
    Strategy: rebuild from original using opcodes and accept_mask toggle.
    """
    o = original.split()
    c = corrected.split()
    matcher = SequenceMatcher(None, o, c)
    out = []
    mask_idx = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            out.extend(o[i1:i2])
        else:
            accept = accept_mask[mask_idx] if mask_idx < len(accept_mask) else False
            mask_idx += 1
            if tag in ("replace", "insert"):
                if accept:
                    out.extend(c[j1:j2])
                else:
                    out.extend(o[i1:i2])  # keep original
            elif tag == "delete":
                if accept:
                    # delete -> skip original tokens
                    pass
                else:
                    out.extend(o[i1:i2])
    return " ".join(out)

# ---------- Export / Import ----------

def export_text(text: str, fmt: str = "txt"):
    """
    Export corrected text as .txt or .docx (BytesIO for Streamlit download).
    """
    if fmt == "txt":
        return text
    if fmt == "docx":
        if Document is None:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        buf = BytesIO()
        doc = Document()
        for para in text.split("\n"):
            doc.add_paragraph(para)
        doc.save(buf)
        buf.seek(0)
        return buf
    raise ValueError("Unsupported format. Use 'txt' or 'docx'.")

def load_text_from_file(uploaded_file) -> str:
    """
    Read .txt or .docx Streamlit UploadedFile and return plain text.
    """
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        if Document is None:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        from docx import Document as Doc
        f = BytesIO(uploaded_file.read())
        doc = Doc(f)
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("Unsupported file type. Upload .txt or .docx")

# ---------- Dictionaries ----------

BUILTIN_DICTIONARIES = {
    "Medical": {"hypertension", "myocardial", "diastolic", "metformin", "angioplasty"},
    "Legal": {"tort", "estoppel", "habeas", "jurisprudence", "fiduciary"},
    "Coding": {"NumPy", "PyTorch", "TensorFlow", "async", "await", "TypeScript"},
}

def build_preserve_terms(
    user_terms: List[str],
    dictionary_tags: List[str],
    uploaded_glossary: str = "",
) -> List[str]:
    terms = set(t.strip() for t in (user_terms or []) if t.strip())
    for tag in dictionary_tags or []:
        terms |= BUILTIN_DICTIONARIES.get(tag, set())
    if uploaded_glossary:
        for line in uploaded_glossary.splitlines():
            if line.strip():
                terms.add(line.strip())
    return sorted(terms)

# ---------- Speech (Optional) ----------

def tts_gtts(text: str):
    """
    Generate an MP3 (BytesIO) using gTTS if available.
    """
    if gTTS is None:
        return None, "gTTS not installed. Run: pip install gTTS"
    buf = BytesIO()
    try:
        gTTS(text).write_to_fp(buf)
        buf.seek(0)
        return buf, None
    except Exception as e:
        return None, f"TTS failed: {e}"
